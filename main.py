import docx

# Taxonomy: Document objects -> Paragraph objects -> Run objects
d = docx.Document('demo.docx')
print(d.paragraphs[0])
p = d.paragraphs[0]
print(p.text)
print(p.runs)

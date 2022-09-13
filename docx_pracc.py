import docx

# Taxonomy: Document objects -> Paragraph objects -> Run objects
# Reading docx
d = docx.Document('demo.docx')
print(d.paragraphs[0])
p = d.paragraphs[0]
print(p.text)
print(p.runs)

# Creating docx
d2 = docx.Document()
d2.add_paragraph('This is a paragraph.')
d2.add_paragraph('And this is another paragraph.')
# Set the first paragraph's first run to bold
d2.paragraphs[0].runs[0].bold = True
# Add another run to the first paragraph and make it italic
d2.paragraphs[0].add_run('\nThis is a new italic run.')
d2.paragraphs[0].runs[1].italic = True
d2.save('demo2.docx')


# Get all the text from a Word document

def get_text(filename):
    doc = docx.Document(filename)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)


print(get_text('demo.docx'))